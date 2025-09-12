import os
import logging
import tempfile
from telegram import Update, ReplyKeyboardMarkup, ReplyKeyboardRemove
from telegram.ext import (
    Updater, CommandHandler, MessageHandler, 
    Filters, ConversationHandler, CallbackContext
)
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# تمكين التسجيل
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)

# مراحل المحادثة
(
    START_CHOICE, NAME, PHONE, EMAIL, ADDRESS, 
    CAREER_OBJECTIVE, EDUCATION, EXPERIENCE, 
    SKILLS, LANGUAGES, ADD_CUSTOM_SECTION, 
    CUSTOM_SECTION_NAME, CUSTOM_SECTION_CONTENT, 
    TEMPLATE, REVIEW, PAYMENT
) = range(16)

# بيانات المستخدم
user_data = {}
cv_file_path = None

# معلومات الدفع
BANK_INFO = """
✅ للدفع عبر البنك:
- اسم المستفيد: عمر محمد السهلي
- البنك: الراجحي  
- رقم الحساب: SA0080000000000000000000
- المبلغ: 25 ريال سعودي

بعد التحويل، أرسل 'تم الدفع' لاستلام السيرة الذاتية.
"""

# أزرار تفاعلية
def create_keyboard(options):
    return ReplyKeyboardMarkup([[option] for option in options], one_time_keyboard=True, resize_keyboard=True)

def start(update, context):
    global cv_file_path
    user_data.clear()
    cv_file_path = None
    
    welcome_msg = (
        "🎯 **مرحباً بك في بوت السيرة الذاتية الاحترافية!**\n\n"
        "سأساعدك في إنشاء سيرة ذاتية إنجليزية احترافية بتصميم عصري.\n\n"
        "💰 **سعر الخدمة: 25 ريال سعودي**\n\n"
        "🚀 **اختر طريقة البدء:**"
    )
    
    update.message.reply_text(welcome_msg, reply_markup=create_keyboard(['📝 بدء إنشاء السيرة', 'ℹ️ معلومات عن البوت']))
    return START_CHOICE

def start_choice(update, context):
    choice = update.message.text
    
    if choice == '📝 بدء إنشاء السيرة':
        update.message.reply_text(
            "👤 **ما هو اسمك بالكامل؟**\n\n"
            "اكتب اسمك كما تريد ظهوره في السيرة الذاتية",
            reply_markup=create_keyboard(['رجوع'])
        )
        return NAME
        
    elif choice == 'ℹ️ معلومات عن البوت':
        info_msg = (
            "🤖 **معلومات عن البوت:**\n\n"
            "• إنشاء سيرة ذاتية إنجليزية احترافية\n"
            "• تصميم ATS-friendly للتوافق مع أنظمة التوظيف\n"
            "• 3 قوالب مختلفة للاختيار\n"
            "• إضافة أقسام مخصصة\n"
            "• حفظ البيانات خلال الجلسة\n"
            "• إمكانية الرجوع والتعديل\n\n"
            "💰 **السعر: 25 ريال سعودي** لكل سيرة ذاتية\n\n"
            "🎯 **للبَدء، اختر 'بدء إنشاء السيرة'**"
        )
        update.message.reply_text(info_msg, reply_markup=create_keyboard(['📝 بدء إنشاء السيرة', 'رجوع']))
        return START_CHOICE
        
    else:
        update.message.reply_text("❌ اختر من الخيارات المتاحة")
        return START_CHOICE

def get_name(update, context):
    if update.message.text.lower() == 'رجوع':
        update.message.reply_text("🔙 عدنا للقائمة الرئيسية:", reply_markup=create_keyboard(['📝 بدء إنشاء السيرة', 'ℹ️ معلومات عن البوت']))
        return START_CHOICE
        
    user_data['name'] = update.message.text
    
    next_msg = (
        "📱 **أدخل رقم جوالك:**\n"
        "مثال: 0512345678"
    )
    update.message.reply_text(next_msg, reply_markup=create_keyboard(['رجوع']))
    return PHONE

def get_phone(update, context):
    if update.message.text.lower() == 'رجوع':
        update.message.reply_text("🔙 عدنا لسؤال الاسم:\nما هو اسمك بالكامل؟")
        return NAME
        
    user_data['phone'] = update.message.text
    
    next_msg = (
        "📧 **أدخل بريدك الإلكتروني:**\n"
        "مثال: name@example.com"
    )
    update.message.reply_text(next_msg, reply_markup=create_keyboard(['رجوع']))
    return EMAIL

def get_email(update, context):
    if update.message.text.lower() == 'رجوع':
        update.message.reply_text("🔙 عدنا لسؤال الجوال:\nأدخل رقم جوالك:")
        return PHONE
        
    user_data['email'] = update.message.text
    
    next_msg = (
        "🏠 **أدخل عنوانك:**\n"
        "مثال: Medina, Saudi Arabia"
    )
    update.message.reply_text(next_msg, reply_markup=create_keyboard(['رجوع', 'تخطي']))
    return ADDRESS

def get_address(update, context):
    if update.message.text.lower() == 'رجوع':
        user_data.pop('email', None)
        update.message.reply_text("🔙 عدنا لسؤال الإيميل:\nأدخل بريدك الإلكتروني:")
        return EMAIL
    elif update.message.text.lower() == 'تخطي':
        user_data['address'] = "Medina, Saudi Arabia"
        update.message.reply_text("✅ تم استخدام عنوان افتراضي.")
    else:
        user_data['address'] = update.message.text
    
    objective_msg = (
        "🎯 **أدخل الهدف المهني (Career Objective):**\n\n"
        "💡 **مثال:**\n"
        "To leverage my technical and sales expertise in building AI-powered digital solutions and driving revenue growth within a forward-thinking organization."
    )
    update.message.reply_text(objective_msg, reply_markup=create_keyboard(['رجوع', 'تخطي']))
    return CAREER_OBJECTIVE

def get_career_objective(update, context):
    if update.message.text.lower() == 'رجوع':
        user_data.pop('address', None)
        update.message.reply_text("🔙 عدنا لسؤال العنوان:\nأدخل عنوانك:")
        return ADDRESS
    elif update.message.text.lower() == 'تخطي':
        user_data['career_objective'] = "Seeking a challenging position to utilize my skills and contribute to organizational growth."
        update.message.reply_text("✅ تم استخدام هدف افتراضي.")
    else:
        user_data['career_objective'] = update.message.text
    
    edu_msg = (
        "🎓 **أدخل مؤهلاتك التعليمية:**\n\n"
        "💡 **مثال:**\n"
        "High School Diploma - Government School - 2011"
    )
    update.message.reply_text(edu_msg, reply_markup=create_keyboard(['رجوع', 'تخطي']))
    return EDUCATION

def get_education(update, context):
    if update.message.text.lower() == 'رجوع':
        user_data.pop('career_objective', None)
        update.message.reply_text("🔙 عدنا لسؤال الهدف المهني:\nأدخل الهدف المهني:")
        return CAREER_OBJECTIVE
    elif update.message.text.lower() == 'تخطي':
        user_data['education'] = "No formal education specified"
        update.message.reply_text("✅ تم تخطي التعليم.")
    else:
        user_data['education'] = update.message.text
    
    exp_msg = (
        "💼 **أدخل خبراتك العملية:**\n\n"
        "💡 **مثال:**\n"
        "Sales Officer | Wahat Al Munawara\n"
        "Jan 2019 – Present\n"
        "• Generated over 300,000 SAR in annual sales\n"
        "• Built long-term client relationships\n\n"
        "Owner & Founder | Digital Developer Establishment\n"
        "Apr 2017 – Jan 2019\n"
        "• Launched and managed smartphone retail business"
    )
    update.message.reply_text(exp_msg, reply_markup=create_keyboard(['رجوع', 'تخطي']))
    return EXPERIENCE

def get_experience(update, context):
    if update.message.text.lower() == 'رجوع':
        user_data.pop('education', None)
        update.message.reply_text("🔙 عدنا لسؤال التعليم:\nأدخل مؤهلاتك التعليمية:")
        return EDUCATION
    elif update.message.text.lower() == 'تخطي':
        user_data['experience'] = "No work experience specified"
        update.message.reply_text("✅ تم تخطي الخبرات.")
    else:
        user_data['experience'] = update.message.text
    
    skills_msg = (
        "🛠️ **أدخل مهاراتك (افصل بينها بفواصل):**\n\n"
        "💡 **مثال:**\n"
        "Sales Strategy, Digital Marketing, CRM, Project Management, Microsoft Office"
    )
    update.message.reply_text(skills_msg, reply_markup=create_keyboard(['رجوع', 'تخطي']))
    return SKILLS

def get_skills(update, context):
    if update.message.text.lower() == 'رجوع':
        user_data.pop('experience', None)
        update.message.reply_text("🔙 عدنا لسؤال الخبرات:\nأدخل خبراتك العملية:")
        return EXPERIENCE
    elif update.message.text.lower() == 'تخطي':
        user_data['skills'] = "No skills specified"
        update.message.reply_text("✅ تم تخطي المهارات.")
    else:
        user_data['skills'] = update.message.text
    
    lang_msg = (
        "🌐 **أدخل اللغات التي تتقنها:**\n\n"
        "💡 **مثال:**\n"
        "Arabic (Native), English (Fluent)"
    )
    update.message.reply_text(lang_msg, reply_markup=create_keyboard(['رجوع', 'تخطي']))
    return LANGUAGES

def get_languages(update, context):
    if update.message.text.lower() == 'رجوع':
        user_data.pop('skills', None)
        update.message.reply_text("🔙 عدنا لسؤال المهارات:\nأدخل مهاراتك:")
        return SKILLS
    elif update.message.text.lower() == 'تخطي':
        user_data['languages'] = "No languages specified"
        update.message.reply_text("✅ تم تخطي اللغات.")
    else:
        user_data['languages'] = update.message.text
    
    # السؤال عن الأقسام الإضافية
    custom_msg = (
        "➕ **هل ترغب في إضافة قسم إضافي؟**\n\n"
        "يمكنك إضافة أقسام مثل:\n"
        "• المشاريع الشخصية\n"
        "• الدورات التدريبية\n"
        "• الجوائز والتكريمات\n"
        "• الهوايات والاهتمامات\n"
        "• المراجع\n\n"
        "اختر 'نعم' لإضافة قسم أو 'تخطي' للمتابعة"
    )
    update.message.reply_text(custom_msg, reply_markup=create_keyboard(['نعم', 'تخطي', 'رجوع']))
    return ADD_CUSTOM_SECTION

def add_custom_section(update, context):
    choice = update.message.text.lower()
    
    if choice == 'رجوع':
        user_data.pop('languages', None)
        update.message.reply_text("🔙 عدنا لسؤال اللغات:\nأدخل اللغات التي تتقنها:")
        return LANGUAGES
    elif choice == 'تخطي':
        # تهيئة الأقسام المخصصة كقائمة فارغة إذا لم تكن موجودة
        if 'custom_sections' not in user_data:
            user_data['custom_sections'] = []
        
        # اختيار القالب
        template_msg = (
            "🎨 **اختر تصميم السيرة الذاتية:**\n\n"
            "1. **كلاسيكي** - تنسيق تقليدي ومهني\n"
            "2. **حديث** - تصميم ATS عصري (موصى به)\n"
            "3. **مبدع** - تصميم أنيق مع خطوط مميزة\n\n"
            "أختر رقم القالب (1, 2, 3):"
        )
        update.message.reply_text(template_msg, reply_markup=create_keyboard(['1', '2', '3', 'رجوع']))
        return TEMPLATE
    elif choice == 'نعم':
        update.message.reply_text(
            "📝 **ما هو اسم القسم الذي تريد إضافته؟**\n\n"
            "مثال: Projects, Certifications, Awards, etc.",
            reply_markup=create_keyboard(['رجوع'])
        )
        return CUSTOM_SECTION_NAME
    else:
        update.message.reply_text("❌ اختر من الخيارات المتاحة")
        return ADD_CUSTOM_SECTION

def get_custom_section_name(update, context):
    if update.message.text.lower() == 'رجوع':
        update.message.reply_text(
            "➕ **هل ترغب في إضافة قسم إضافي؟",
            reply_markup=create_keyboard(['نعم', 'تخطي', 'رجوع'])
        )
        return ADD_CUSTOM_SECTION
        
    # تخزين اسم القسم مؤقتاً
    context.user_data['current_section_name'] = update.message.text
    
    update.message.reply_text(
        f"📋 **أدخل محتوى قسم '{update.message.text}':**\n\n"
        "اكتب المعلومات التي تريد إضافتها في هذا القسم",
        reply_markup=create_keyboard(['رجوع'])
    )
    return CUSTOM_SECTION_CONTENT

def get_custom_section_content(update, context):
    if update.message.text.lower() == 'رجوع':
        update.message.reply_text(
            "📝 **ما هو اسم القسم الذي تريد إضافته؟**",
            reply_markup=create_keyboard(['رجوع'])
        )
        return CUSTOM_SECTION_NAME
        
    # الحصول على اسم القسم من البيانات المؤقتة
    section_name = context.user_data.get('current_section_name', 'Custom Section')
    
    # إضافة القسم إلى بيانات المستخدم
    if 'custom_sections' not in user_data:
        user_data['custom_sections'] = []
    
    user_data['custom_sections'].append({
        'name': section_name,
        'content': update.message.text
    })
    
    # تنظيف البيانات المؤقتة
    context.user_data.pop('current_section_name', None)
    
    # السؤال عن إضافة المزيد من الأقسام
    update.message.reply_text(
        f"✅ تم إضافة قسم '{section_name}' بنجاح!\n\n"
        "هل ترغب في إضافة قسم آخر؟",
        reply_markup=create_keyboard(['نعم', 'لا', 'رجوع'])
    )
    return ADD_CUSTOM_SECTION

def choose_template(update, context):
    if update.message.text.lower() == 'رجوع':
        # إذا كان هناك أقسام مخصصة، احذفها وارجع
        if user_data.get('custom_sections'):
            user_data.pop('custom_sections', None)
        update.message.reply_text(
            "➕ **هل ترغب في إضافة قسم إضافي؟",
            reply_markup=create_keyboard(['نعم', 'تخطي', 'رجوع'])
        )
        return ADD_CUSTOM_SECTION
        
    template_choice = update.message.text
    templates = {
        '1': 'classic',
        '2': 'modern', 
        '3': 'creative'
    }
    
    if template_choice in templates:
        user_data['template'] = templates[template_choice]
        
        # معاينة البيانات - السطر المصحح
        preview_msg = (
            "📋 **لمحة عن بياناتك:**\n\n"
            f"👤 **الاسم:** {user_data.get('name', 'N/A')}\n"
            f"📞 **الجوال:** {user_data.get('phone', 'N/A')}\n"
            f"📧 **الإيميل:** {user_data.get('email', 'N/A')}\n"
            f"🏠 **العنوان:** {user_data.get('address', 'N/A')}\n"
            f"🎯 **الهدف:** {user_data.get('career_objective', 'N/A')[:50]}...\n"
            f"🎓 **التعليم:** {user_data.get('education', 'N/A')[:50]}...\n"
            f"💼 **الخبرات:** {user_data.get('experience', 'N/A')[:50]}...\n"
            f"🛠️ **المهارات:** {user_data.get('skills', 'N/A')[:50]}...\n"
            f"🌐 **اللغات:** {user_data.get('languages', 'N/A')}\n"
        )
        
        # إضافة الأقسام المخصصة للمعاينة إذا وجدت - السطر المصحح
        if user_data.get('custom_sections'):
            preview_msg += f"➕ **الأقسام الإضافية:** {len(user_data['custom_sections'])} قسم\n"
        
        preview_msg += f"🎨 **التصميم:** {user_data.get('template', 'N/A')}\n\n"
        preview_msg += "هل تريد المتابعة وإنشاء السيرة الذاتية؟"
        
        update.message.reply_text(preview_msg, reply_markup=create_keyboard(['نعم', 'لا', 'تعديل']))
        return REVIEW
    else:
        update.message.reply_text("❌ اختر رقم صحيح (1, 2, 3)")
        return TEMPLATE

def review_data(update, context):
    choice = update.message.text.lower()
    
    if choice == 'نعم':
        try:
            global cv_file_path
            update.message.reply_text("⏳ جاري إنشاء سيرتك الذاتية...")
            cv_file_path = create_professional_cv(user_data, user_data.get('template', 'modern'))
            
            success_msg = (
                f"✅ **تهانينا {user_data.get('name')}!**\n\n"
                "تم إنشاء سيرتك الذاتية بنجاح 🎉\n\n"
                "💰 **السعر: 25 ريال سعودي**\n\n"
                f"{BANK_INFO}\n\n"
                "أرسل 'تم الدفع' بعد التحويل لاستلام الملف."
            )
            update.message.reply_text(success_msg, reply_markup=create_keyboard(['تم الدفع']))
            return PAYMENT
            
        except Exception as e:
            logger.error(f"CV creation error: {e}")
            update.message.reply_text("❌ حدث خطأ في الإنشاء. حاول /start مرة أخرى.")
            return ConversationHandler.END
            
    elif choice == 'تعديل':
        # إنشاء قائمة خيارات التعديل
        options = [
            'الاسم', 'الجوال', 'الإيميل', 'العنوان', 'الهدف', 
            'التعليم', 'الخبرات', 'المهارات', 'اللغات', 'التصميم'
        ]
        
        # إضافة خيارات للأقسام المخصصة إذا وجدت
        if user_data.get('custom_sections'):
            for i, section in enumerate(user_data['custom_sections']):
                options.append(f'القسم {i+1}: {section["name"]}')
        
        update.message.reply_text("🔧 اختر ما تريد تعديله:", reply_markup=create_keyboard(options))
        return REVIEW
        
    else:  # لا أو أي رد آخر
        update.message.reply_text("❌ تم إلغاء العملية. اكتب /start للبدء من جديد.")
        return ConversationHandler.END

def check_payment(update, context):
    if "تم الدفع" in update.message.text.lower():
        try:
            if cv_file_path and os.path.exists(cv_file_path):
                with open(cv_file_path, 'rb') as doc_file:
                    update.message.reply_document(
                        document=doc_file,
                        filename=f"CV_{user_data.get('name', 'User')}.docx",
                        caption="🎉 **ها هي سيرتك الذاتية الجاهزة!**\n\nشكراً لثقتك بنا 🌟"
                    )
                update.message.reply_text("✅ تم الإرسال بنجاح! اكتب /start لإنشاء سيرة جديدة.")
            else:
                update.message.reply_text("❌ لم يتم العثور على الملف. اكتب /start للبدء من جديد.")
        except Exception as e:
            logger.error(f"File send error: {e}")
            update.message.reply_text("❌ خطأ في الإرسال. حاول /start مرة أخرى.")
        return ConversationHandler.END
    else:
        update.message.reply_text("⚠️ أرسل 'تم الدفع' بعد اكتمال التحويل.")
        return PAYMENT

def cancel(update, context):
    update.message.reply_text(
        "❌ تم إلغاء العملية.\n\n"
        "اكتب /start عندما تكون جاهزاً للبدء 🚀",
        reply_markup=ReplyKeyboardRemove()
    )
    return ConversationHandler.END

def create_professional_cv(data, template_name):
    try:
        temp_dir = tempfile.gettempdir()
        cv_filename = f"CV_{data.get('name', 'User').replace(' ', '_')}.docx"
        cv_path = os.path.join(temp_dir, cv_filename)
        
        doc = Document()
        
        # تطبيق القالب المختار
        if template_name == 'classic':
            apply_classic_template(doc, data)
        elif template_name == 'modern':
            apply_modern_ats_template(doc, data)
        elif template_name == 'creative':
            apply_creative_template(doc, data)
        else:
            apply_modern_ats_template(doc, data)  # افتراضي
        
        doc.save(cv_path)
        logger.info(f"CV created with {template_name} template: {cv_path}")
        return cv_path
        
    except Exception as e:
        logger.error(f"CV creation error: {e}")
        raise

def apply_modern_ats_template(doc, data):
    """التصميم الحديث ATS-Friendly"""
    # === الإعداد العام ===
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # === الاسم ===
    name = doc.add_paragraph()
    name_run = name.add_run(data.get('name', '').upper())
    name_run.font.size = Pt(16)
    name_run.bold = True
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.paragraph_format.space_after = Pt(6)
    
    # === معلومات الاتصال ===
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run(f"Phone: {data.get('phone', '')} | ")
    contact.add_run(f"Email: {data.get('email', '')} | ")
    contact.add_run(f"Address: {data.get('address', 'Medina, Saudi Arabia')}")
    contact.paragraph_format.space_after = Pt(12)
    
    # === الهدف المهني ===
    if data.get('career_objective'):
        doc.add_heading('CAREER OBJECTIVE', level=1)
        objective = doc.add_paragraph(data.get('career_objective'))
        objective.paragraph_format.space_after = Pt(12)
    
    # === الخبرات ===
    if data.get('experience') and data.get('experience') != "No work experience specified":
        doc.add_heading('EXPERIENCE', level=1)
        experience = doc.add_paragraph(data.get('experience'))
        experience.paragraph_format.space_after = Pt(12)
    
    # === المهارات ===
    if data.get('skills') and data.get('skills') != "No skills specified":
        doc.add_heading('SKILLS', level=1)
        skills = doc.add_paragraph(data.get('skills'))
        skills.paragraph_format.space_after = Pt(12)
    
    # === التعليم ===
    if data.get('education') and data.get('education') != "No formal education specified":
        doc.add_heading('EDUCATION', level=1)
        education = doc.add_paragraph(data.get('education'))
        education.paragraph_format.space_after = Pt(12)
    
    # === اللغات ===
    if data.get('languages') and data.get('languages') != "No languages specified":
        doc.add_heading('LANGUAGES', level=1)
        languages = doc.add_paragraph(data.get('languages'))
        languages.paragraph_format.space_after = Pt(12)
    
    # === الأقسام المخصصة ===
    if data.get('custom_sections'):
        for section in data['custom_sections']:
            doc.add_heading(section['name'].upper(), level=1)
            content = doc.add_paragraph(section['content'])
            content.paragraph_format.space_after = Pt(12)

def apply_classic_template(doc, data):
    """القوالب الكلاسيكي"""
    doc.add_heading('CURRICULUM VITAE', 0)
    add_personal_info_simple(doc, data)
    add_section_simple(doc, 'CAREER OBJECTIVE', data.get('career_objective'))
    add_section_simple(doc, 'EXPERIENCE', data.get('experience'))
    add_section_simple(doc, 'SKILLS', data.get('skills'))
    add_section_simple(doc, 'EDUCATION', data.get('education'))
    add_section_simple(doc, 'LANGUAGES', data.get('languages'))
    
    # الأقسام المخصصة
    if data.get('custom_sections'):
        for section in data['custom_sections']:
            add_section_simple(doc, section['name'].upper(), section['content'])

def apply_creative_template(doc, data):
    """القوالب الإبداعي"""
    title = doc.add_heading('CURRICULUM VITAE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style.font.size = Pt(18)
    title.style.font.name = 'Georgia'
    
    add_personal_info_simple(doc, data)
    add_section_simple(doc, 'CAREER OBJECTIVE', data.get('career_objective'))
    add_section_simple(doc, 'PROFESSIONAL EXPERIENCE', data.get('experience'))
    add_section_simple(doc, 'SKILLS & COMPETENCIES', data.get('skills'))
    add_section_simple(doc, 'EDUCATION', data.get('education'))
    add_section_simple(doc, 'LANGUAGES', data.get('languages'))
    
    # الأقسام المخصصة
    if data.get('custom_sections'):
        for section in data['custom_sections']:
            add_section_simple(doc, section['name'].upper(), section['content'])

def add_personal_info_simple(doc, data):
    """معلومات شخصية مبسطة"""
    doc.add_heading('PERSONAL INFORMATION', level=1)
    p = doc.add_paragraph()
    p.add_run('Name: ').bold = True
    p.add_run(data.get('name', 'N/A'))
    p.add_run('\nPhone: ').bold = True
    p.add_run(data.get('phone', 'N/A'))
    p.add_run('\nEmail: ').bold = True
    p.add_run(data.get('email', 'N/A'))
    p.add_run('\nAddress: ').bold = True
    p.add_run(data.get('address', 'N/A'))
    p.paragraph_format.space_after = Pt(12)

def add_section_simple(doc, title, content):
    """إضافة قسم مبسط"""
    if content and "No " not in content:
        doc.add_heading(title, level=1)
        doc.add_paragraph(content)
        doc.add_paragraph().paragraph_format.space_after = Pt(12)

def error_handler(update, context):
    logger.error(f'Bot error: {context.error}')
    if update and update.message:
        update.message.reply_text(
            "❌ حدث خطأ غير متوقع.\n\n"
            "اكتب /start للمحاولة مرة أخرى 🔄"
        )

def main():
    try:
        token = os.getenv('TELEGRAM_BOT_TOKEN')
        if not token:
            logger.error("❌ TELEGRAM_BOT_TOKEN not set")
            return
        
        updater = Updater(token, use_context=True)
        dp = updater.dispatcher
        
        dp.add_error_handler(error_handler)
        
        conv_handler = ConversationHandler(
            entry_points=[CommandHandler('start', start), MessageHandler(Filters.text & ~Filters.command, start)],
            states={
                START_CHOICE: [MessageHandler(Filters.text & ~Filters.command, start_choice)],
                NAME: [MessageHandler(Filters.text & ~Filters.command, get_name)],
                PHONE: [MessageHandler(Filters.text & ~Filters.command, get_phone)],
                EMAIL: [MessageHandler(Filters.text & ~Filters.command, get_email)],
                ADDRESS: [MessageHandler(Filters.text & ~Filters.command, get_address)],
                CAREER_OBJECTIVE: [MessageHandler(Filters.text & ~Filters.command, get_career_objective)],
                EDUCATION: [MessageHandler(Filters.text & ~Filters.command, get_education)],
                EXPERIENCE: [MessageHandler(Filters.text & ~Filters.command, get_experience)],
                SKILLS: [MessageHandler(Filters.text & ~Filters.command, get_skills)],
                LANGUAGES: [MessageHandler(Filters.text & ~Filters.command, get_languages)],
                ADD_CUSTOM_SECTION: [MessageHandler(Filters.text & ~Filters.command, add_custom_section)],
                CUSTOM_SECTION_NAME: [MessageHandler(Filters.text & ~Filters.command, get_custom_section_name)],
                CUSTOM_SECTION_CONTENT: [MessageHandler(Filters.text & ~Filters.command, get_custom_section_content)],
                TEMPLATE: [MessageHandler(Filters.text & ~Filters.command, choose_template)],
                REVIEW: [MessageHandler(Filters.text & ~Filters.command, review_data)],
                PAYMENT: [MessageHandler(Filters.text & ~Filters.command, check_payment)],
            },
            fallbacks=[CommandHandler('cancel', cancel)],
        )
        
        dp.add_handler(conv_handler)
        
        # بدء البوت
        updater.start_polling()
        logger.info("✅ Bot is running with enhanced features!")
        updater.idle()
        
    except Exception as e:
        logger.error(f"❌ Bot startup error: {e}")

if __name__ == '__main__':
    main()
